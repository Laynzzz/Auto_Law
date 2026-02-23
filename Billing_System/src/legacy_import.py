"""Legacy invoice importer — parse old monthly invoice .docx files and
import case data into the master dataset.

Supports the legacy format with 4 tables:
    Table 0: Header (firm info, invoice #, date)
    Table 1: Recipient (To, re, P.O. #)
    Table 2: Line items (date, DESCRIPTION, HOURS, RATE, AMOUNT)
    Table 3: Footer

The DESCRIPTION column format is:
    INDEX_NUMBER [CASE_CAPTION]
e.g. "LT-306306-Q-LT Fake abc vs. Jake Banner, Tom Banner"
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document

from src.dataset import upsert_row


# Pattern: index number is the leading alphanumeric-dash token(s)
# e.g. "LT-306306-Q-LT", "12345/2026"
_IDX_PATTERN = re.compile(r"^(\S+)\s*(.*)")


def parse_legacy_invoice(file_path: str | Path) -> list[dict]:
    """Parse a legacy monthly invoice .docx and return a list of case dicts.

    Each dict has keys: appearance_date, index_number, case_caption, charge_amount.
    """
    doc = Document(str(file_path))

    if len(doc.tables) < 3:
        raise ValueError(
            f"Expected at least 3 tables in legacy invoice, found {len(doc.tables)}: {file_path}"
        )

    table = doc.tables[2]  # line items table
    cases: list[dict] = []

    for i, row in enumerate(table.rows):
        if i == 0:
            continue  # skip header row

        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 5:
            continue

        date_str, desc, hours, rate, amount = cells[:5]

        # Skip empty rows and total row
        if not date_str and not desc:
            continue
        if "TOTAL" in desc.upper():
            continue

        # Parse date (formats: "2/11/26", "02/11/2026", etc.)
        if not date_str:
            continue
        try:
            dt = datetime.strptime(date_str, "%m/%d/%y")
        except ValueError:
            try:
                dt = datetime.strptime(date_str, "%m/%d/%Y")
            except ValueError:
                continue  # skip rows with unparseable dates

        app_date = dt.strftime("%Y-%m-%d")

        # Parse description: index_number + optional case caption
        m = _IDX_PATTERN.match(desc)
        if m:
            index_number = m.group(1)
            case_caption = m.group(2).strip() if m.group(2).strip() else index_number
        else:
            index_number = desc
            case_caption = desc

        # Parse amount (0 if empty)
        try:
            charge = float(amount.replace(",", "").replace("$", "")) if amount else 0.0
        except (ValueError, AttributeError):
            charge = 0.0

        cases.append({
            "appearance_date": app_date,
            "index_number": index_number,
            "case_caption": case_caption,
            "charge_amount": charge,
        })

    return cases


def import_legacy_invoice(
    firm_name: str,
    file_path: str | Path,
) -> list[tuple[str, dict]]:
    """Parse a legacy invoice and import cases into the firm's dataset.

    Returns list of (action, case_dict) where action is "inserted" or "updated".
    Holds a single firm lock for the entire batch to avoid per-row lock overhead.
    """
    from src.file_lock import FirmFileLock

    cases = parse_legacy_invoice(file_path)
    results: list[tuple[str, dict]] = []

    with FirmFileLock(firm_name):
        for case in cases:
            action = upsert_row(firm_name, case, _hold_lock=False)
            results.append((action, case))

    return results
