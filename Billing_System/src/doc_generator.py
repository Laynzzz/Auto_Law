"""Document generator — fill Word template placeholders and convert to PDF.

Uses template/perdiem.docx as the single template for per-case invoices.
Replaces [[placeholder]] tokens in runs, preserving formatting.
Converts the filled .docx to .pdf via docx2pdf (Word COM on Windows).

Output structure:
    invoice/{FirmName}/{YYYY}/{Mon}/Week of MM-DD-YYYY/report/word/MM-DD-YYYY Case Name.docx
    invoice/{FirmName}/{YYYY}/{Mon}/Week of MM-DD-YYYY/report/pdf/MM-DD-YYYY Case Name.pdf
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx2pdf import convert

from src.config import get_firm, load_config
from src.dataset import find_row_by_key, week_range, PROJECT_ROOT

TEMPLATE_PATH = PROJECT_ROOT / "template" / "perdiem.docx"


# ── Date formatting ──────────────────────────────────────────────────

def _ordinal(day: int) -> str:
    """Return day with ordinal suffix: 1st, 2nd, 3rd, 4th, ..."""
    if 11 <= day <= 13:
        return f"{day}th"
    suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return f"{day}{suffix}"


def _format_date_display(date_str: str) -> str:
    """Convert YYYY-MM-DD or datetime to 'February 20th, 2026'."""
    if isinstance(date_str, datetime):
        dt = date_str
    else:
        dt = datetime.strptime(str(date_str).split(" ")[0], "%Y-%m-%d")
    return f"{dt.strftime('%B')} {_ordinal(dt.day)}, {dt.year}"


def _parse_date(date_str) -> datetime:
    """Parse a date string or datetime into a datetime object."""
    if isinstance(date_str, datetime):
        return date_str
    return datetime.strptime(str(date_str).split(" ")[0], "%Y-%m-%d")


# ── Placeholder map ──────────────────────────────────────────────────

def _build_placeholder_map(case: dict, firm: dict) -> dict[str, str]:
    """Build a {placeholder: value} map from case data and firm config."""
    amount = case.get("charge_amount")
    if amount is not None:
        fee_str = f"{float(amount):,.2f}"
    else:
        fee_str = "0.00"

    date_display = _format_date_display(case.get("appearance_date") or "")

    return {
        "[[invoice number]]": str(case.get("invoice_number") or ""),
        "[[Date]]": date_display,
        "[[Name]]": firm.get("contact_name", ""),
        "[[Company Name]]": firm.get("name", ""),
        "[[Address 1]]": firm.get("address_1", ""),
        "[[Address 2]]": firm.get("address_2", ""),
        "[[Case Name]]": str(case.get("case_caption") or ""),
        "[[Index Number]]": str(case.get("index_number") or ""),
        "[[Location]]": str(case.get("court") or ""),
        "[[Result]]": str(case.get("outcome") or ""),
        "[[Fee]]": fee_str,
    }


# ── Fill template ────────────────────────────────────────────────────

def _replace_in_runs(paragraph, placeholders: dict[str, str]) -> None:
    """Replace [[placeholder]] tokens in a paragraph's runs."""
    for run in paragraph.runs:
        for token, value in placeholders.items():
            if token in run.text:
                run.text = run.text.replace(token, value)


def fill_template(case: dict, firm: dict, output_docx: Path) -> Path:
    """Fill the perdiem.docx template with case data and save to output_docx."""
    doc = Document(str(TEMPLATE_PATH))
    placeholders = _build_placeholder_map(case, firm)

    for paragraph in doc.paragraphs:
        _replace_in_runs(paragraph, placeholders)

    # Also check tables (the template has a "PER DIEM" header table)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_runs(paragraph, placeholders)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return output_docx


# ── PDF conversion ───────────────────────────────────────────────────

def convert_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    """Convert a .docx file to .pdf using Word COM. Returns the PDF path."""
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    convert(str(docx_path), str(pdf_path))
    return pdf_path


# ── Full pipeline ────────────────────────────────────────────────────

def generate_invoice(
    firm_name: str,
    index_number: str,
    appearance_date: str,
    config: dict | None = None,
    keep_docx: bool = False,
) -> Path:
    """Full pipeline: look up case, fill template, convert to PDF.

    Output structure:
        invoice/{firm}/{YYYY}/{Mon}/Week of MM-DD-YYYY/report/word/MM-DD-YYYY Case Name.docx
        invoice/{firm}/{YYYY}/{Mon}/Week of MM-DD-YYYY/report/pdf/MM-DD-YYYY Case Name.pdf

    Returns the path to the generated PDF.
    """
    if config is None:
        config = load_config()

    firm = get_firm(firm_name, config)
    case = find_row_by_key(firm_name, index_number, appearance_date)

    if case is None:
        raise ValueError(
            f"Case not found: firm={firm_name}, "
            f"index={index_number}, date={appearance_date}"
        )

    inv_num = case.get("invoice_number")
    if not inv_num:
        raise ValueError(
            f"Case has no invoice number. Run 'assign-invoices' first.\n"
            f"  firm={firm_name}, index={index_number}, date={appearance_date}"
        )

    # Build output paths
    dt = _parse_date(case.get("appearance_date"))
    year_folder = str(dt.year)                     # e.g. "2026"
    month_folder = dt.strftime("%b")               # e.g. "Feb"
    date_prefix = dt.strftime("%m-%d-%Y")          # e.g. "02-20-2026"
    caption = str(case.get("case_caption") or "case")
    filename = f"{date_prefix} {caption}"

    # Week folder: "Week of MM-DD-YYYY" based on Monday of the case's week
    monday, _ = week_range(dt.date() if isinstance(dt, datetime) else dt)
    week_folder = f"Week of {monday.strftime('%m-%d-%Y')}"

    base_dir = (
        PROJECT_ROOT / "invoice" / firm_name
        / year_folder / month_folder / week_folder / "report"
    )
    docx_out = base_dir / "word" / f"{filename}.docx"
    pdf_out = base_dir / "pdf" / f"{filename}.pdf"

    # Fill template → word
    fill_template(case, firm, docx_out)

    # Convert → pdf
    convert_to_pdf(docx_out, pdf_out)

    # Clean up intermediate .docx unless requested
    if not keep_docx and docx_out.exists():
        docx_out.unlink()
        # Remove word dir if empty
        word_dir = docx_out.parent
        if word_dir.exists() and not any(word_dir.iterdir()):
            word_dir.rmdir()

    return pdf_out
