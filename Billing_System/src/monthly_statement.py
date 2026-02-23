"""Monthly Statement of Account generator.

Uses template/monthly_statement.docx as the template.
Replaces [[placeholder]] tokens (handling cross-run splits) and fills the
case table with data rows.

Output structure:
    invoice/{FirmName}/{YYYY}/{Mon}/Monthly Statement {Mon} {YYYY}.docx
    invoice/{FirmName}/{YYYY}/{Mon}/Monthly Statement {Mon} {YYYY}.pdf
"""

import copy
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx2pdf import convert

from src.config import get_firm, load_config
from src.dataset import PROJECT_ROOT, get_data_root, month_range, query_by_date_range, _to_date
from src.doc_generator import _format_date_display
from src.weekly_statement import _replace_in_paragraph, _set_cell_text, _clone_row, _clear_row

TEMPLATE_PATH = PROJECT_ROOT / "template" / "monthly_statement.docx"


# ── Fill template ────────────────────────────────────────────────────

def _fill_monthly_template(
    firm_name: str,
    year: int,
    month: int,
    cases: list[dict],
    firm: dict,
    output_docx: Path,
) -> Path:
    """Fill the monthly_statement.docx template and save."""
    doc = Document(str(TEMPLATE_PATH))

    month_name = date(year, month, 1).strftime("%B")  # e.g. "February"

    # Period string: "02/01/26 - 02/28/26"
    first_day, last_day = month_range(year, month)
    period_str = f"{first_day.strftime('%m/%d/%y')} - {last_day.strftime('%m/%d/%y')}"

    # Header/body placeholders
    header_placeholders = {
        "[[week date]]": period_str,
        "[[Date]]": _format_date_display(date.today().isoformat()),
        "[[Name]]": firm.get("contact_name", ""),
        "[[Company Name]]": firm.get("name", ""),
        "[[Address 1]]": firm.get("address_1", ""),
        "[[Address 2]]": firm.get("address_2", ""),
    }

    # Replace header/body placeholders
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, header_placeholders)

    # Also replace in table[0] (the "MONTHLY STATMENT" banner)
    if doc.tables:
        for row in doc.tables[0].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, header_placeholders)

    # ── Fill case table (table[1]) ───────────────────────────────────
    case_table = doc.tables[1]
    num_cols = 4  # Date, Index No., Case Caption, Amount

    # Template has: row 0 = header, row 1 = template data row,
    # rows 2-25 = empty, row 26 = total row.
    pre_allocated = 25  # rows 1..25
    total_row_idx = 26  # original index of the TOTAL row

    # If we need more rows than pre-allocated, clone the template row
    if len(cases) > pre_allocated:
        extra_needed = len(cases) - pre_allocated
        for _ in range(extra_needed):
            _clone_row(case_table, 1)
        total_row_idx = 1 + len(cases)

    # Fill case data into rows 1..N
    total_fee = 0.0
    for i, case in enumerate(cases):
        row = case_table.rows[1 + i]
        d = _to_date(case.get("appearance_date"))
        date_str = d.strftime("%m/%d/%Y") if d else ""
        amt = float(case.get("charge_amount") or 0)
        total_fee += amt

        _set_cell_text(row, 0, date_str)
        _set_cell_text(row, 1, str(case.get("index_number") or ""))
        _set_cell_text(row, 2, str(case.get("case_caption") or ""))
        _set_cell_text(row, 3, f"${amt:,.2f}")

    # Clear any unused pre-allocated rows
    used_data_rows = len(cases)
    total_data_slots = max(pre_allocated, len(cases))
    for i in range(used_data_rows, total_data_slots):
        row = case_table.rows[1 + i]
        _clear_row(row, num_cols)

    # Replace [[total fee]] in the total row
    total_row = case_table.rows[total_row_idx]
    for cell in total_row.cells:
        for paragraph in cell.paragraphs:
            _replace_in_paragraph(paragraph, {
                "[[total fee]]": f"{total_fee:,.2f}",
            })

    # Clear placeholder from template row 1 if no cases
    if not cases:
        row1 = case_table.rows[1]
        for placeholder in ["[[case date]]", "[[case no]]", "[[case caption]]", "[[case fee]]"]:
            for cell in row1.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, {placeholder: ""})

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return output_docx


# ── Full pipeline ────────────────────────────────────────────────────

def generate_monthly_statement(
    firm_name: str,
    year: int,
    month: int,
    config: dict | None = None,
    keep_docx: bool = False,
) -> Path:
    """Generate a monthly statement PDF for a firm.

    Returns the path to the generated PDF.
    """
    if config is None:
        config = load_config()

    firm = get_firm(firm_name, config)
    cases = query_by_date_range(
        firm_name, *month_range(year, month)
    )

    # Output paths — flat in the month folder
    month_abbr = date(year, month, 1).strftime("%b")
    filename = f"Monthly Statement {month_abbr} {year}"

    base_dir = (
        get_data_root() / "invoice" / firm_name
        / str(year) / month_abbr
    )
    docx_out = base_dir / f"{filename}.docx"
    pdf_out = base_dir / f"{filename}.pdf"

    # Fill template
    _fill_monthly_template(firm_name, year, month, cases, firm, docx_out)

    # Convert to PDF
    pdf_out.parent.mkdir(parents=True, exist_ok=True)
    convert(str(docx_out), str(pdf_out))

    # Clean up intermediate .docx (don't remove month folder)
    if not keep_docx and docx_out.exists():
        docx_out.unlink()

    return pdf_out
