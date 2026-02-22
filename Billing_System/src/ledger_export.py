"""Master Ledger Export — full-history ledger per firm.

Generates a PDF (and optional XLSX) showing every case for a firm,
with totals and aging analysis for outstanding invoices.

Output structure:
    invoice/{FirmName}/ledger/Ledger as of MM-DD-YYYY.pdf
    invoice/{FirmName}/ledger/Ledger as of MM-DD-YYYY.xlsx  (optional)
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx2pdf import convert
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, numbers
from openpyxl.utils import get_column_letter

from src.config import load_config
from src.dataset import PROJECT_ROOT, load_dataset, _to_date
from src.doc_generator import _ordinal


# ── Aging buckets ────────────────────────────────────────────────────

AGING_BRACKETS = [
    ("0-30 days", 0, 30),
    ("31-60 days", 31, 60),
    ("61-90 days", 61, 90),
    ("90+ days", 91, None),
]


def _compute_aging(cases: list[dict], as_of: date) -> list[tuple[str, int, float]]:
    """Return list of (bracket_label, count, total) for unpaid cases."""
    buckets: dict[str, tuple[int, float]] = {
        label: (0, 0.0) for label, _, _ in AGING_BRACKETS
    }

    for case in cases:
        paid = str(case.get("paid_status") or "").strip()
        if paid == "Paid":
            continue
        d = _to_date(case.get("appearance_date"))
        if d is None:
            continue
        age = (as_of - d).days
        amt = float(case.get("charge_amount") or 0)

        for label, lo, hi in AGING_BRACKETS:
            if hi is None:
                if age >= lo:
                    count, total = buckets[label]
                    buckets[label] = (count + 1, total + amt)
                    break
            elif lo <= age <= hi:
                count, total = buckets[label]
                buckets[label] = (count + 1, total + amt)
                break

    return [(label, *buckets[label]) for label, _, _ in AGING_BRACKETS]


# ── PDF (Word → PDF) ────────────────────────────────────────────────

TABLE_COLUMNS = [
    "Date", "Invoice #", "Index #", "Case Caption",
    "Court", "Status", "Amount", "Paid", "Payment Date",
]
COL_WIDTHS = [0.78, 0.78, 0.88, 1.45, 0.95, 0.65, 0.65, 0.5, 0.78]


def _build_ledger_doc(
    firm_name: str,
    as_of: date,
    cases: list[dict],
    output_docx: Path,
) -> Path:
    """Create the ledger .docx and save it."""
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("PICERNO & ASSOCIATES, PLLC")
    run.bold = True
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Master Ledger")
    run.font.size = Pt(12)

    # Firm & as-of date
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run("Firm: ").bold = True
    info.add_run(firm_name)

    as_of_p = doc.add_paragraph()
    as_of_p.add_run("As of: ").bold = True
    month_name = as_of.strftime("%B")
    as_of_p.add_run(f"{month_name} {_ordinal(as_of.day)}, {as_of.year}")

    doc.add_paragraph()

    # Case table
    table = doc.add_table(rows=1, cols=len(TABLE_COLUMNS))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, name in enumerate(TABLE_COLUMNS):
        cell = table.rows[0].cells[i]
        cell.text = name
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)

    for i, w in enumerate(COL_WIDTHS):
        table.columns[i].width = Inches(w)

    total_billed = 0.0
    total_paid = 0.0

    for case in cases:
        row = table.add_row()
        d = _to_date(case.get("appearance_date"))
        date_str = d.strftime("%m/%d/%Y") if d else ""
        amt = float(case.get("charge_amount") or 0)
        total_billed += amt

        paid_status = str(case.get("paid_status") or "").strip()
        if paid_status == "Paid":
            total_paid += amt
            paid_display = "Paid"
        elif paid_status == "Partial":
            paid_display = "Partial"
        else:
            paid_display = ""

        pay_date = _to_date(case.get("payment_date"))
        pay_date_str = pay_date.strftime("%m/%d/%Y") if pay_date else ""

        values = [
            date_str,
            str(case.get("invoice_number") or ""),
            str(case.get("index_number") or ""),
            str(case.get("case_caption") or ""),
            str(case.get("court") or ""),
            str(case.get("case_status") or ""),
            f"${amt:,.2f}",
            paid_display,
            pay_date_str,
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = val
            for p in cell.paragraphs:
                if i == 6:  # Amount
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif i in (7, 8):  # Paid, Payment Date
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(8)

    # Total row
    total_row = table.add_row()
    for i in range(len(TABLE_COLUMNS)):
        cell = total_row.cells[i]
        if i == 5:
            cell.text = "Total:"
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(8)
        elif i == 6:
            cell.text = f"${total_billed:,.2f}"
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(8)

    # Summary
    doc.add_paragraph()
    outstanding = total_billed - total_paid

    for label, value, bold in [
        ("Total Cases", str(len(cases)), False),
        ("Total Billed", f"${total_billed:,.2f}", False),
        ("Total Paid", f"${total_paid:,.2f}", False),
        ("Outstanding", f"${outstanding:,.2f}", True),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: {value}")
        run.font.size = Pt(10)
        if bold:
            run.bold = True

    # Aging analysis
    aging = _compute_aging(cases, as_of)
    has_aging = any(count > 0 for _, count, _ in aging)

    if has_aging:
        doc.add_paragraph()
        aging_h = doc.add_paragraph()
        run = aging_h.add_run("Aging Analysis (Outstanding)")
        run.bold = True
        run.font.size = Pt(11)

        aging_table = doc.add_table(rows=1, cols=3)
        aging_table.style = "Light Grid Accent 1"
        aging_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, name in enumerate(["Period", "Cases", "Amount"]):
            cell = aging_table.rows[0].cells[i]
            cell.text = name
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        aging_table.columns[0].width = Inches(1.5)
        aging_table.columns[1].width = Inches(1.0)
        aging_table.columns[2].width = Inches(1.5)

        for label, count, total in aging:
            row = aging_table.add_row()
            for i, val in enumerate([label, str(count), f"${total:,.2f}"]):
                cell = row.cells[i]
                cell.text = val
                for p in cell.paragraphs:
                    p.alignment = (
                        WD_ALIGN_PARAGRAPH.RIGHT if i == 2
                        else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    for r in p.runs:
                        r.font.size = Pt(9)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return output_docx


# ── XLSX export ──────────────────────────────────────────────────────

XLSX_COLUMNS = [
    "Appearance Date", "Invoice #", "Index #", "Case Caption",
    "Court", "Outcome", "Case Status", "Charge Amount",
    "Invoice Sent Date", "Paid Status", "Payment Date", "Notes",
]

HEADER_FILL = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=10)


def _build_ledger_xlsx(
    firm_name: str,
    as_of: date,
    cases: list[dict],
    output_xlsx: Path,
) -> Path:
    """Create the ledger .xlsx and save it."""
    from src.dataset import COLUMNS  # dataset column keys

    wb = Workbook()
    ws = wb.active
    ws.title = "Ledger"

    # Header row
    for col_idx, name in enumerate(XLSX_COLUMNS, start=1):
        cell = ws.cell(row=1, column=col_idx, value=name)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")

    # Data rows
    for row_idx, case in enumerate(cases, start=2):
        for col_idx, col_key in enumerate(COLUMNS, start=1):
            val = case.get(col_key)
            # Format dates as strings for readability
            d = _to_date(val) if col_key.endswith("_date") else None
            if d is not None:
                val = d.strftime("%m/%d/%Y")
            elif col_key == "charge_amount" and val is not None:
                val = float(val)
            ws.cell(row=row_idx, column=col_idx, value=val)

    # Format charge_amount column as currency
    amt_col = COLUMNS.index("charge_amount") + 1
    for row_idx in range(2, len(cases) + 2):
        ws.cell(row=row_idx, column=amt_col).number_format = '$#,##0.00'

    # Auto-width columns
    for col_idx in range(1, len(XLSX_COLUMNS) + 1):
        letter = get_column_letter(col_idx)
        max_len = len(XLSX_COLUMNS[col_idx - 1])
        for row_idx in range(2, min(len(cases) + 2, 50)):
            val = ws.cell(row=row_idx, column=col_idx).value
            if val is not None:
                max_len = max(max_len, len(str(val)))
        ws.column_dimensions[letter].width = min(max_len + 3, 40)

    # Summary rows below data
    summary_row = len(cases) + 3
    ws.cell(row=summary_row, column=1, value="Summary").font = Font(bold=True, size=11)

    total_billed = sum(float(c.get("charge_amount") or 0) for c in cases)
    total_paid = sum(
        float(c.get("charge_amount") or 0) for c in cases
        if str(c.get("paid_status") or "").strip() == "Paid"
    )
    outstanding = total_billed - total_paid

    for i, (label, value) in enumerate([
        ("Total Cases", len(cases)),
        ("Total Billed", total_billed),
        ("Total Paid", total_paid),
        ("Outstanding", outstanding),
    ]):
        r = summary_row + 1 + i
        ws.cell(row=r, column=1, value=label).font = Font(bold=True)
        cell = ws.cell(row=r, column=2, value=value)
        if isinstance(value, float):
            cell.number_format = '$#,##0.00'

    # Aging analysis
    aging = _compute_aging(cases, as_of)
    aging_row = summary_row + 6
    ws.cell(row=aging_row, column=1, value="Aging Analysis (Outstanding)").font = Font(bold=True, size=11)
    ws.cell(row=aging_row + 1, column=1, value="Period").font = Font(bold=True)
    ws.cell(row=aging_row + 1, column=2, value="Cases").font = Font(bold=True)
    ws.cell(row=aging_row + 1, column=3, value="Amount").font = Font(bold=True)

    for i, (label, count, total) in enumerate(aging):
        r = aging_row + 2 + i
        ws.cell(row=r, column=1, value=label)
        ws.cell(row=r, column=2, value=count)
        cell = ws.cell(row=r, column=3, value=total)
        cell.number_format = '$#,##0.00'

    ws.freeze_panes = "A2"

    output_xlsx.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output_xlsx)
    return output_xlsx


# ── Full pipeline ────────────────────────────────────────────────────

def export_ledger(
    firm_name: str,
    as_of: date | None = None,
    config: dict | None = None,
    keep_docx: bool = False,
    xlsx: bool = True,
) -> dict[str, Path]:
    """Export a firm's master ledger as PDF (and optional XLSX).

    Returns dict with keys 'pdf' and optionally 'xlsx'.
    """
    if config is None:
        config = load_config()
    if as_of is None:
        as_of = date.today()

    cases = load_dataset(firm_name)
    # Sort by appearance_date
    cases.sort(key=lambda r: _to_date(r.get("appearance_date")) or date.min)

    date_prefix = as_of.strftime("%m-%d-%Y")
    filename = f"Ledger as of {date_prefix}"

    base_dir = PROJECT_ROOT / "invoice" / firm_name / "ledger"
    docx_out = base_dir / f"{filename}.docx"
    pdf_out = base_dir / f"{filename}.pdf"

    # Build PDF via Word
    _build_ledger_doc(firm_name, as_of, cases, docx_out)
    pdf_out.parent.mkdir(parents=True, exist_ok=True)
    convert(str(docx_out), str(pdf_out))

    if not keep_docx and docx_out.exists():
        docx_out.unlink()

    result: dict[str, Path] = {"pdf": pdf_out}

    # XLSX
    if xlsx:
        xlsx_out = base_dir / f"{filename}.xlsx"
        _build_ledger_xlsx(firm_name, as_of, cases, xlsx_out)
        result["xlsx"] = xlsx_out

    return result
