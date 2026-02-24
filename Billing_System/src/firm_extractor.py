"""Extract law firm contact info from real invoice .docx files.

Parses the "Bill To" block from each firm's invoice to extract:
firm name, address, phone, email, and contact names.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document


# ── Regex patterns ───────────────────────────────────────────────────

_PHONE_RE = re.compile(
    r"""
    (?:p\.?\s*)?                      # optional "p." or "p " prefix
    \(?\d{3}\)?[\s.\-]*\d{3}[\s.\-]*\d{4}  # (XXX) XXX-XXXX and variants
    (?:\s*(?:ext|x)\.?\s*\d+)?        # optional extension
    """,
    re.VERBOSE | re.IGNORECASE,
)

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w.-]+\.\w{2,}", re.IGNORECASE)

_NAME_EMAIL_RE = re.compile(
    r"([A-Z][A-Za-z.\-']+(?:\s+[A-Z][A-Za-z.\-']+)+)\s*[<(]([\w.+-]+@[\w.-]+\.\w{2,})[>)]",
)

_ZIP_RE = re.compile(r"\b\d{5}(?:-\d{4})?\b")

_SUITE_FLOOR_RE = re.compile(
    r"(?:suite|ste|floor|fl|room|rm|#)\s*\S+",
    re.IGNORECASE,
)

_FAX_LINE_RE = re.compile(r"\b(?:fax|facsimile)\b", re.IGNORECASE)


# ── Helpers ──────────────────────────────────────────────────────────


def _looks_like_address(line: str) -> bool:
    """Heuristic: line looks like a street address or city/state/zip."""
    if _ZIP_RE.search(line):
        return True
    # Street-address indicators
    if re.search(r"\d+\s+\w+", line) and not _PHONE_RE.search(line):
        # Has "number word" pattern but is not a phone line
        addr_words = (
            "street", "st", "avenue", "ave", "boulevard", "blvd",
            "road", "rd", "drive", "dr", "lane", "ln", "place", "pl",
            "plaza", "floor", "suite", "ste", "broadway",
        )
        lower = line.lower()
        if any(w in lower for w in addr_words):
            return True
        # Has digits + comma (e.g. "123 Main, Suite 400")
        if "," in line:
            return True
    # "City, State" pattern
    if re.search(r"[A-Za-z]+,\s*[A-Za-z]{2,}\s+\d{5}", line):
        return True
    return False


def _generate_initials(firm_name: str) -> str:
    """Auto-generate initials from firm name (2-3 uppercase letters)."""
    # Remove common suffixes
    cleaned = re.sub(
        r",?\s*(?:P\.?C\.?|LLC|LLP|PLLC|L\.L\.P\.?|Esq\.?|Inc\.?|Corp\.?)\.?\s*$",
        "",
        firm_name,
        flags=re.IGNORECASE,
    ).strip()
    # Remove punctuation for initial extraction
    cleaned = re.sub(r"[^\w\s]", "", cleaned)
    words = [w for w in cleaned.split() if w[0].isupper() or w.isupper()]
    if not words:
        words = cleaned.split()
    if len(words) == 1:
        return words[0][:2].upper()
    # Take first letter of each significant word, max 3
    initials = "".join(w[0].upper() for w in words[:3])
    return initials


def _extract_to_lines(doc: Document) -> list[str] | None:
    """Find the 'To:' block lines from either tables or paragraphs.

    Returns cleaned lines (after 'To:' prefix) or None if not found.
    """
    # Strategy 1: check tables[1].cell(0,0) if doc has >= 2 tables
    if len(doc.tables) >= 2:
        cell_text = doc.tables[1].cell(0, 0).text.strip()
        if cell_text.upper().startswith("TO"):
            return _split_cell_text(cell_text)

    # Strategy 2: for docs with tables[0] that might have To: in cell
    if len(doc.tables) >= 1:
        cell_text = doc.tables[0].cell(0, 0).text.strip()
        if cell_text.upper().startswith("TO"):
            return _split_cell_text(cell_text)

    # Strategy 3: scan paragraphs for "TO:" and collect until "RE:"
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    to_idx = None
    for i, text in enumerate(paragraphs):
        if text.upper().startswith("TO:") or text.upper() == "TO":
            to_idx = i
            break

    if to_idx is None:
        return None

    lines: list[str] = []
    # The "TO:" paragraph itself might have content after "TO:"
    first_line = re.sub(r"^TO:?\s*", "", paragraphs[to_idx], flags=re.IGNORECASE).strip()
    if first_line:
        lines.append(first_line)

    # Collect subsequent paragraphs until "RE:" or end
    for j in range(to_idx + 1, min(to_idx + 15, len(paragraphs))):
        text = paragraphs[j].strip()
        if text.upper().startswith("RE:") or text.upper() == "RE":
            break
        # Strip trailing "RE:" or "RE" that got concatenated (no space before)
        text = re.sub(r"RE:?\s*$", "", text).strip()
        if text:
            lines.append(text)

    return lines if lines else None


def _split_cell_text(cell_text: str) -> list[str]:
    """Split a table cell's text into lines, stripping the 'To:' prefix."""
    raw = cell_text.replace("\r", "\n").split("\n")
    lines: list[str] = []
    for line in raw:
        line = line.strip()
        if not line:
            continue
        # Strip "To:" prefix from first non-empty line
        if not lines and re.match(r"^TO:?\s*", line, re.IGNORECASE):
            line = re.sub(r"^TO:?\s*", "", line, flags=re.IGNORECASE).strip()
            if not line:
                continue
        lines.append(line)
    return lines


# ── Main extraction ──────────────────────────────────────────────────


def extract_firm_info(docx_path: Path) -> dict:
    """Parse a .docx invoice and return extracted firm metadata.

    Returns a dict with keys: firm_name, initials, contact_name,
    address_1, address_2, phone, billing_email, cc_emails.
    """
    doc = Document(str(docx_path))
    lines = _extract_to_lines(doc)

    result = {
        "firm_name": "",
        "initials": "",
        "contact_name": "",
        "address_1": "",
        "address_2": "",
        "phone": "",
        "billing_email": "",
        "cc_emails": [],
    }

    if not lines:
        return result

    # ── Parse lines ──────────────────────────────────────────────
    # Line 1 is always the firm name
    result["firm_name"] = lines[0]
    result["initials"] = _generate_initials(lines[0])

    addresses: list[str] = []
    phones: list[str] = []
    emails: list[str] = []
    contact_names: list[str] = []

    for line in lines[1:]:
        # Skip fax lines
        if _FAX_LINE_RE.search(line) and not _PHONE_RE.search(line):
            continue

        # Check for "Name <email>" pattern
        name_email_match = _NAME_EMAIL_RE.search(line)
        if name_email_match:
            contact_names.append(name_email_match.group(1).strip())
            emails.append(name_email_match.group(2).strip())
            continue

        # Check for standalone email
        email_match = _EMAIL_RE.search(line)
        if email_match and not _looks_like_address(line):
            emails.append(email_match.group())
            # If there's text before the email, might be a contact name
            prefix = line[:email_match.start()].strip().rstrip(":")
            if prefix and not _PHONE_RE.search(prefix):
                contact_names.append(prefix)
            continue

        # Check for phone number
        phone_match = _PHONE_RE.search(line)
        if phone_match:
            # Check if this line is ONLY a phone (not address with phone)
            if not _looks_like_address(line):
                phones.append(phone_match.group().strip())
                # Line might also have a fax after — ignore fax
                continue

        # Remaining lines: address or contact name
        if _looks_like_address(line):
            addresses.append(line)
        elif _ZIP_RE.search(line):
            addresses.append(line)
        else:
            # Could be a second contact name line, or dept name
            # If it's a short non-address line after firm name, treat as
            # part of address (e.g. "Attn: Someone", dept, floor)
            if len(addresses) < 2:
                addresses.append(line)
            # else skip

    # ── Assign parsed values ─────────────────────────────────────
    if addresses:
        result["address_1"] = addresses[0]
    if len(addresses) >= 2:
        result["address_2"] = addresses[1]

    if phones:
        result["phone"] = phones[0]

    if emails:
        result["billing_email"] = emails[0]
        if len(emails) > 1:
            result["cc_emails"] = emails[1:]

    if contact_names:
        result["contact_name"] = contact_names[0]

    return result


# ── Scanning ─────────────────────────────────────────────────────────


def scan_all_firms(invoices_dir: Path) -> tuple[list[dict], list[str]]:
    """Scan all firm folders and extract metadata from one .docx per firm.

    Args:
        invoices_dir: Path to the invoices root (e.g. ``invoice/2026 Invoices``).

    Returns:
        (firms, warnings) — list of extracted firm dicts and list of warning strings.
    """
    firms: list[dict] = []
    warnings: list[str] = []

    if not invoices_dir.is_dir():
        warnings.append(f"Directory not found: {invoices_dir}")
        return firms, warnings

    for folder in sorted(invoices_dir.iterdir()):
        if not folder.is_dir():
            continue

        # Find a .docx file: prefer top-level, fall back to Paid/ subfolder
        docx_file = _pick_docx(folder)

        if docx_file is None:
            warnings.append(f"No .docx found: {folder.name} (PDF-only or empty)")
            continue

        try:
            info = extract_firm_info(docx_file)
        except Exception as exc:
            warnings.append(f"Error reading {folder.name}/{docx_file.name}: {exc}")
            continue

        info["folder_name"] = folder.name
        info["source_file"] = docx_file.name

        # If extraction yielded no firm name, use folder name
        if not info["firm_name"]:
            info["firm_name"] = folder.name
            info["initials"] = _generate_initials(folder.name)
            warnings.append(
                f"Could not parse firm name from {folder.name}/{docx_file.name}; "
                "using folder name"
            )

        firms.append(info)

    return firms, warnings


def _pick_docx(folder: Path) -> Path | None:
    """Pick the first .docx in *folder*, falling back to Paid/ subfolder."""
    # Top-level .docx files
    candidates = sorted(folder.glob("*.docx"))
    if candidates:
        return candidates[0]

    # Fall back to Paid/ subfolder
    paid_dir = folder / "Paid"
    if paid_dir.is_dir():
        candidates = sorted(paid_dir.glob("*.docx"))
        if candidates:
            return candidates[0]

    return None
