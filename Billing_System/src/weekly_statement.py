"""Weekly Statement of Account generator.

Uses template/weekly_statement.docx as the template.
Replaces [[placeholder]] tokens (handling cross-run splits) and fills the
case table with data rows.

Output structure:
    invoice/{FirmName}/{YYYY}/{Mon}/Week of MM-DD-YYYY/Week of MM-DD-YYYY.docx
    invoice/{FirmName}/{YYYY}/{Mon}/Week of MM-DD-YYYY/Week of MM-DD-YYYY.pdf
"""

import copy
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx2pdf import convert

from src.config import get_firm, load_config
from src.dataset import PROJECT_ROOT, get_data_root, query_by_date_range, week_range, _to_date
from src.doc_generator import _format_date_display

TEMPLATE_PATH = PROJECT_ROOT / "template" / "weekly_statement.docx"


# ── Cross-run placeholder replacement ────────────────────────────────

def _replace_in_paragraph(paragraph, placeholders: dict[str, str]) -> None:
    """Replace [[placeholder]] tokens that may span multiple runs.

    Joins all run texts, performs replacements, then redistributes text
    back into runs (all text goes into the first run; remaining runs are
    cleared).  This preserves the first run's formatting.
    """
    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(r.text for r in runs)

    # Quick check — skip if no placeholder present
    if "[[" not in full_text:
        return

    new_text = full_text
    for token, value in placeholders.items():
        new_text = new_text.replace(token, value)

    if new_text == full_text:
        return  # nothing changed

    # Put all text in the first run, clear the rest
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def _replace_all_placeholders(doc: Document, placeholders: dict[str, str]) -> None:
    """Replace placeholders in all paragraphs and table cells."""
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, placeholders)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, placeholders)


# ── Table row helpers ────────────────────────────────────────────────

def _clone_row(table, source_row_idx: int) -> None:
    """Clone a table row (deep copy of XML) and insert it right after the source row."""
    source_tr = table.rows[source_row_idx]._tr
    new_tr = copy.deepcopy(source_tr)
    source_tr.addnext(new_tr)


def _set_cell_text(row, col_idx: int, text: str) -> None:
    """Set cell text preserving existing paragraph/run formatting.

    If the cell has existing runs, reuses the first run's formatting.
    If empty (no runs), creates a run matching the template style
    (Calibri 10pt).
    """
    cell = row.cells[col_idx]
    for paragraph in cell.paragraphs:
        runs = paragraph.runs
        if runs:
            runs[0].text = text
            for r in runs[1:]:
                r.text = ""
        else:
            run = paragraph.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.bold = False


def _clear_row(row, num_cols: int) -> None:
    """Clear all cells in a row."""
    for i in range(num_cols):
        _set_cell_text(row, i, "")


# ── Fill template ────────────────────────────────────────────────────

def _fill_weekly_template(
    firm_name: str,
    monday: date,
    friday: date,
    cases: list[dict],
    firm: dict,
    output_docx: Path,
) -> Path:
    """Fill the weekly_statement.docx template and save."""
    doc = Document(str(TEMPLATE_PATH))

    # Period string: "02/16/26 - 02/20/26"
    period_str = f"{monday.strftime('%m/%d/%y')} - {friday.strftime('%m/%d/%y')}"

    # Header/body placeholders
    header_placeholders = {
        "[[week date]]": period_str,
        "[[Date]]": _format_date_display(date.today().isoformat()),
        "[[Name]]": firm.get("contact_name", ""),
        "[[Company Name]]": firm.get("name", ""),
        "[[Address 1]]": firm.get("address_1", ""),
        "[[Address 2]]": firm.get("address_2", ""),
    }

    # Replace header/body placeholders (paragraphs + table[0] title)
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, header_placeholders)

    # Also replace in table[0] (the "WEEKLY STATMENT" banner) in case it has placeholders
    if doc.tables:
        for row in doc.tables[0].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, header_placeholders)

    # ── Fill case table (table[1]) ───────────────────────────────────
    case_table = doc.tables[1]
    num_cols = 4  # Date, Index No., Case Caption, Amount

    # Template has: row 0 = header, row 1 = template data row,
    # rows 2-25 = empty, row 26 = total row.
    # Pre-allocated data slots = rows 1 through 25 (25 rows).
    pre_allocated = 25  # rows 1..25
    total_row_idx = 26  # original index of the TOTAL row

    # If we need more rows than pre-allocated, clone the template row
    if len(cases) > pre_allocated:
        extra_needed = len(cases) - pre_allocated
        # Insert clones before the total row (after last pre-allocated row)
        for _ in range(extra_needed):
            _clone_row(case_table, 1)  # clone template row format
        # Total row index shifts
        total_row_idx = 1 + len(cases)

    # Fill case data into rows 1..N
    total_fee = 0.0
    for i, case in enumerate(cases):
        row = case_table.rows[1 + i]
        d = _to_date(case.get("appearance_date"))
        date_str = d.strftime("%m/%d/%Y") if d else ""
        amt = float(case.get("charge_amount") or 0)
        total_fee += amt

        _set_cell_text(row, 0, date_str)
        _set_cell_text(row, 1, str(case.get("index_number") or ""))
        _set_cell_text(row, 2, str(case.get("case_caption") or ""))
        _set_cell_text(row, 3, f"${amt:,.2f}")

    # Clear any unused pre-allocated rows (if fewer cases than slots)
    used_data_rows = len(cases)
    total_data_slots = max(pre_allocated, len(cases))
    for i in range(used_data_rows, total_data_slots):
        row = case_table.rows[1 + i]
        _clear_row(row, num_cols)

    # Replace [[total fee]] in the total row
    total_row = case_table.rows[total_row_idx]
    for cell in total_row.cells:
        for paragraph in cell.paragraphs:
            _replace_in_paragraph(paragraph, {
                "[[total fee]]": f"{total_fee:,.2f}",
            })

    # Also clear placeholder from template row 1 if no cases
    if not cases:
        row1 = case_table.rows[1]
        for placeholder in ["[[case date]]", "[[case no]]", "[[case caption]]", "[[case fee]]"]:
            for cell in row1.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, {placeholder: ""})

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return output_docx


# ── Full pipeline ────────────────────────────────────────────────────

def generate_weekly_statement(
    firm_name: str,
    week_of: date,
    config: dict | None = None,
    keep_docx: bool = False,
) -> Path:
    """Generate a weekly statement PDF for a firm.

    week_of: any date within the desired week (Mon-Fri range is computed).
    Returns the path to the generated PDF.
    """
    if config is None:
        config = load_config()

    firm = get_firm(firm_name, config)
    monday, friday = week_range(week_of)
    cases = query_by_date_range(firm_name, monday, friday)

    # Output paths — flat in the week folder (no subfolders)
    date_prefix = monday.strftime("%m-%d-%Y")
    week_folder = f"Week of {date_prefix}"

    base_dir = (
        get_data_root() / "invoice" / firm_name
        / str(monday.year) / monday.strftime("%b") / week_folder
    )
    docx_out = base_dir / f"{week_folder}.docx"
    pdf_out = base_dir / f"{week_folder}.pdf"

    # Fill template
    _fill_weekly_template(firm_name, monday, friday, cases, firm, docx_out)

    # Convert to PDF
    pdf_out.parent.mkdir(parents=True, exist_ok=True)
    convert(str(docx_out), str(pdf_out))

    # Clean up intermediate .docx (don't remove week folder — it has other files)
    if not keep_docx and docx_out.exists():
        docx_out.unlink()

    return pdf_out
